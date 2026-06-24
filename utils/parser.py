"""
utils/parser.py
---------------
Shared resume parser for both Method 1 (TF-IDF) and Method 2 (BERT).
Handles PDF and DOCX files.
"""

import os
import re
import pdfplumber
from docx import Document


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"  [WARNING] Could not parse PDF: {filepath} — {e}")
    return text.strip()


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text.strip() + "\n"
        # Also extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text.strip() + " "
                text += "\n"
    except Exception as e:
        print(f"  [WARNING] Could not parse DOCX: {filepath} — {e}")
    return text.strip()


def clean_text(text: str) -> str:
    """
    Basic text cleaning:
    - Remove excessive whitespace and newlines
    - Remove special characters (keep letters, numbers, basic punctuation)
    - Lowercase
    """
    # Replace multiple newlines/spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    # Remove non-ASCII characters (optional — comment out if multilingual)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    # Remove special characters except common punctuation
    text = re.sub(r'[^a-zA-Z0-9\s\.\,\-\+\/\@]', ' ', text)
    # Lowercase
    text = text.lower().strip()
    return text


def load_all_resumes(folder_path: str, clean: bool = True) -> list[dict]:
    """
    Load all PDF and DOCX resumes from a folder.

    Returns:
        List of dicts:
        [
            {"filename": "john_doe.pdf", "raw_text": "...", "cleaned_text": "..."},
            ...
        ]
    """
    supported = ('.pdf', '.docx')
    resumes = []

    files = [f for f in os.listdir(folder_path) if f.lower().endswith(supported)]

    if not files:
        print(f"[ERROR] No PDF or DOCX files found in: {folder_path}")
        return resumes

    print(f"[INFO] Found {len(files)} resume files. Parsing...\n")

    for i, filename in enumerate(files, 1):
        filepath = os.path.join(folder_path, filename)
        ext = os.path.splitext(filename)[1].lower()

        if ext == '.pdf':
            raw_text = extract_text_from_pdf(filepath)
        elif ext == '.docx':
            raw_text = extract_text_from_docx(filepath)
        else:
            continue

        cleaned = clean_text(raw_text) if clean else raw_text

        if not cleaned:
            print(f"  [SKIP] Empty content after parsing: {filename}")
            continue

        resumes.append({
            "filename": filename,
            "raw_text": raw_text,
            "cleaned_text": cleaned
        })

        if i % 20 == 0 or i == len(files):
            print(f"  Parsed {i}/{len(files)} resumes...")

    print(f"\n[DONE] Successfully loaded {len(resumes)} resumes.\n")
    return resumes


def load_job_description(jd_text: str, clean: bool = True) -> str:
    """
    Accepts a raw JD string (paste from notebook), cleans and returns it.
    """
    if clean:
        return clean_text(jd_text)
    return jd_text.strip()
